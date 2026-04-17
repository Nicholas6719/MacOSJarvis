"""
Voice-driven file management for Jarvis.

Search via macOS mdfind (Spotlight CLI) — no third-party search.
Preview via the orb WebView at localhost:3000 OR native Quick Look,
selected by FILE_PREVIEW_MODE.

Public surface:
    FILE_PREVIEW_MODE                       "orb" | "quicklook"
    search_file(query)                      -> list[str]
    get_file_type(filepath)                 -> "image"|"pdf"|"text"|"docx"|"other"
    prepare_preview(filepath)               -> dict (see docstring)
    cleanup_temp_preview(filepath)          -> None
    move_file(source, destination)          -> (bool, str)
    rename_file(filepath, new_name)         -> (bool, str)
    resolve_destination(spoken_location)    -> str
"""

import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional


# ── Feature flag ──────────────────────────────────────────────────────────────
# "orb"       → serve preview through localhost:3000 and render in WebView.
# "quicklook" → skip orb; trigger native macOS Quick Look via qlmanage.
# This is the ONLY toggle point. No other code needs to change.
FILE_PREVIEW_MODE = "orb"


# ── Paths ─────────────────────────────────────────────────────────────────────
_HOME = Path.home()

# The Jarvis project directory itself. Anything inside it — jarvis_memory.db,
# config.json, the Python source, kokoro weights, the frontend build, etc. —
# is OFF-LIMITS to voice file operations. Jarvis must never move, rename,
# describe, or even surface its own internals as a candidate.
#
# Resolved at import time from this file's own location, with a $JARVIS_DATA_DIR
# override for the packaged macOS app (where the code may live inside the
# bundle while data sits in ~/Library/Application Support/Jarvis).
_JARVIS_PROJECT_DIR = Path(__file__).resolve().parent
_JARVIS_DATA_DIR = Path(os.environ.get("JARVIS_DATA_DIR", "")).resolve() \
    if os.environ.get("JARVIS_DATA_DIR") else None

def _protected_roots() -> list[Path]:
    roots = [_JARVIS_PROJECT_DIR]
    if _JARVIS_DATA_DIR and _JARVIS_DATA_DIR != _JARVIS_PROJECT_DIR:
        roots.append(_JARVIS_DATA_DIR)
    return roots


def is_protected_path(filepath: str) -> bool:
    """True if `filepath` lives inside the Jarvis project dir or its data
    dir. Voice file operations must refuse any protected path — we never
    want Jarvis moving, renaming, or describing its own config, memory
    database, model weights, or source code via the voice pipeline."""
    try:
        p = Path(filepath).expanduser().resolve()
    except Exception:
        return False
    for root in _protected_roots():
        try:
            p.relative_to(root)
            return True
        except ValueError:
            continue
    return False


# File extensions that are developer / system artifacts, not user documents.
# Voice file operations skip these even when mdfind matches them — the
# user almost certainly doesn't mean "move my jarvis_memory.db by voice".
_EXCLUDED_EXTENSIONS = {
    ".db", ".sqlite", ".sqlite3", ".sqlite-journal", ".db-journal",
    ".db-shm", ".db-wal", ".sqlite-shm", ".sqlite-wal",
    ".pyc", ".pyo", ".pyd", ".so", ".dylib", ".a", ".o", ".bundle",
    ".log", ".lock", ".pid", ".cache", ".tmp", ".temp", ".bak", ".swp",
    ".class", ".jar", ".gguf", ".onnx", ".bin", ".pt", ".ckpt", ".safetensors",
    ".DS_Store", ".ds_store",
    # macOS app-library internal formats
    ".tvdb", ".musicdb", ".itl", ".itc", ".itdb", ".tagset", ".plist",
}

_COMMON_DESTINATIONS = {
    "desktop":   str(_HOME / "Desktop"),
    "downloads": str(_HOME / "Downloads"),
    "download":  str(_HOME / "Downloads"),
    "documents": str(_HOME / "Documents"),
    "document":  str(_HOME / "Documents"),
    "pictures":  str(_HOME / "Pictures"),
    "picture":   str(_HOME / "Pictures"),
    "photos":    str(_HOME / "Pictures"),
    "music":     str(_HOME / "Music"),
    "movies":    str(_HOME / "Movies"),
    "videos":    str(_HOME / "Movies"),
    "home":      str(_HOME),
}

# Paths we never want to surface to the user — system and hidden locations.
_EXCLUDED_PATH_PREFIXES = (
    "/System/",
    "/Library/",
    "/private/",
    "/usr/",
    "/bin/",
    "/sbin/",
    "/opt/",
    "/Applications/",  # hide bundled app resources from file search
    str(_HOME / "Library") + "/",
)
_EXCLUDED_PATH_FRAGMENTS = (
    "/.Trash/",
    "/.Spotlight-",
    "/.DocumentRevisions-",
    "/.fseventsd/",
    "/node_modules/",
    # macOS app library bundles — these contain internal databases the
    # user should never be moving by voice (Music, TV, Photos, iMovie).
    ".tvlibrary/",
    ".musiclibrary/",
    ".photoslibrary/",
    ".imovielibrary/",
    ".aplibrary/",
    # App bundles and frameworks
    ".app/Contents/",
    ".bundle/Contents/",
    ".framework/",
    # Git and build dirs
    "/.git/",
    "/.venv/",
    "/.venv311/",
    "/__pycache__/",
    "/build/",
    "/dist/",
    "/target/",
)


_TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".rst", ".log", ".csv", ".tsv", ".json",
    ".xml", ".yaml", ".yml", ".ini", ".cfg", ".conf", ".toml",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".html", ".css", ".sh", ".swift",
    ".c", ".cpp", ".h", ".hpp", ".java", ".rb", ".go", ".rs",
}
_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".heic", ".webp", ".tiff"}


# ── Search ────────────────────────────────────────────────────────────────────

def _is_excluded(path: str) -> bool:
    # Jarvis's own files first — never surface them as candidates.
    if is_protected_path(path):
        return True
    for prefix in _EXCLUDED_PATH_PREFIXES:
        if path.startswith(prefix):
            return True
    for frag in _EXCLUDED_PATH_FRAGMENTS:
        if frag in path:
            return True
    # Skip developer/binary extensions — not what the user means by "my file".
    ext = Path(path).suffix.lower()
    if ext in _EXCLUDED_EXTENSIONS:
        return True
    # Anything under a hidden directory in the user's home (e.g. ~/.cache)
    rel = path[len(str(_HOME)):] if path.startswith(str(_HOME)) else path
    for part in rel.split(os.sep):
        if part.startswith(".") and part not in ("", "."):
            return True
    return False


# Words the LLM often drags into the extracted "query" that aren't part of
# the actual filename. We strip them when tokenizing so a messy query like
# "rmv file from my desktop" still finds RMV-RealID-Application-Steps.png.
_QUERY_STOPWORDS = {
    "a", "an", "the", "my", "your", "our", "their", "this", "that", "these",
    "those", "from", "to", "in", "on", "into", "onto", "at", "with", "please",
    "file", "files", "document", "doc", "docs",
    "folder", "directory", "called", "named", "titled", "image", "images",
    "picture", "photo", "and", "or", "of", "for",
    "me", "us", "is", "it", "some", "any", "all", "pdf",
    # Location words — the user is pointing to a folder, not naming
    # the file. Without these, "rmv file from my desktop" would match
    # any path containing "desktop".
    "desktop", "downloads", "download", "documents", "pictures", "photos",
    "music", "movies", "videos", "home", "trash", "icloud",
    # Common action/relational verbs that creep into the extracted
    # query when the LLM includes too much of the utterance.
    "move", "rename", "find", "locate", "summarize", "summarise",
    "describe", "read", "open", "put", "send", "transfer",
    # Jarvis's own name and adjacent words. If these leak into the
    # query (e.g. the LLM carried the wake prefix through, or STT
    # confused RMV with Jarvis), pass-3 would match the project's
    # own jarvis_memory.db. Treat them as filler so Jarvis never
    # searches for itself by voice.
    "jarvis", "jarvisapp", "memory", "config", "settings",
    "database", "log", "logs", "db",
}


def _run_mdfind(args: list[str]) -> list[str]:
    try:
        result = subprocess.run(
            args,
            capture_output=True,
            text=True,
            timeout=8,
        )
    except (subprocess.TimeoutExpired, FileNotFoundError) as e:
        print(f"[file_manager] mdfind failed: {e}")
        return []
    except Exception as e:
        print(f"[file_manager] mdfind error: {e}")
        return []
    return [line.strip() for line in (result.stdout or "").splitlines() if line.strip()]


def _filter_and_rank(raw: list[str]) -> list[str]:
    candidates: list[tuple[float, str]] = []
    for path in raw:
        if _is_excluded(path):
            continue
        try:
            mtime = os.path.getmtime(path)
        except OSError:
            continue
        candidates.append((mtime, path))
    candidates.sort(key=lambda x: x[0], reverse=True)
    return [path for _, path in candidates[:5]]


def _tokenize_query(q: str) -> list[str]:
    """Strip filler words and punctuation, keep tokens ≥2 chars."""
    tokens = [t for t in re.split(r"[^A-Za-z0-9]+", q.lower()) if t]
    return [t for t in tokens if len(t) >= 2 and t not in _QUERY_STOPWORDS]


def search_file(query: str) -> list[str]:
    """Search the Mac by filename via mdfind. Returns up to 5 user-relevant
    matches sorted most-recently-modified first.

    Strategy — try progressively looser queries so messy LLM-extracted
    strings still find real files:
      1. literal substring:   `mdfind -name "<q>"`
      2. AND query on all tokens (stopwords stripped) via kMDItemFSName
      3. per-token search, first hit wins
    """
    q = (query or "").strip()
    if not q:
        return []

    print(f"[file_manager] search_file query={q!r}")

    # Pass 1 — literal substring match on the whole query.
    raw = _run_mdfind(["mdfind", "-name", q])
    ranked = _filter_and_rank(raw)
    if ranked:
        print(f"[file_manager] pass-1 matches={len(ranked)}")
        return ranked

    # Pass 2 — tokenize, keep distinctive words, AND them together.
    tokens = _tokenize_query(q)
    if tokens:
        and_query = " && ".join(f'kMDItemFSName == "*{t}*"cd' for t in tokens)
        raw = _run_mdfind(["mdfind", and_query])
        ranked = _filter_and_rank(raw)
        if ranked:
            print(f"[file_manager] pass-2 tokens={tokens} matches={len(ranked)}")
            return ranked

        # Pass 3 — last resort, try each token alone (most distinctive first,
        # rough heuristic: longer tokens are more distinctive). Take the
        # first token that yields any hits.
        for t in sorted(set(tokens), key=len, reverse=True):
            raw = _run_mdfind(["mdfind", "-name", t])
            ranked = _filter_and_rank(raw)
            if ranked:
                print(f"[file_manager] pass-3 token={t!r} matches={len(ranked)}")
                return ranked

    print("[file_manager] no matches")
    return []


# ── File type ─────────────────────────────────────────────────────────────────

def get_file_type(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()
    if ext in _IMAGE_EXTENSIONS:
        return "image"
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in _TEXT_EXTENSIONS:
        return "text"
    return "other"


# ── Preview ───────────────────────────────────────────────────────────────────
#
# The preview protocol — one file at a time. Python tells ws_server which
# local file to serve through /preview_file, then pushes a WebSocket
# "file_preview" event telling the orb how to render it.
#
# prepare_preview() returns a dict with:
#   file_type:       "image" | "pdf" | "text" | "docx" | "other"
#   serve_url:       str or None   (when a file needs to be fetched)
#   text_content:    str or None   (inline text for text/docx-fallback/other)
#   temp_path:       str or None   (for cleanup_temp_preview)

_PREVIEW_SERVE_URL = "http://localhost:3000/preview_file"


def _read_text_file(path: str, max_bytes: int = 200_000) -> str:
    try:
        with open(path, "rb") as f:
            raw = f.read(max_bytes + 1)
        truncated = len(raw) > max_bytes
        try:
            text = raw[:max_bytes].decode("utf-8")
        except UnicodeDecodeError:
            text = raw[:max_bytes].decode("utf-8", errors="replace")
        if truncated:
            text += "\n\n… (truncated)"
        return text
    except Exception as e:
        return f"(Could not read file: {e})"


def _human_size(n_bytes: int) -> str:
    units = ("B", "KB", "MB", "GB", "TB")
    size = float(n_bytes)
    i = 0
    while size >= 1024 and i < len(units) - 1:
        size /= 1024
        i += 1
    return f"{size:.1f} {units[i]}" if i > 0 else f"{int(size)} {units[i]}"


def _other_card_text(filepath: str) -> str:
    name = Path(filepath).name
    try:
        size = _human_size(os.path.getsize(filepath))
    except OSError:
        size = "unknown size"
    return f"{name}\n{size}\n{filepath}"


def _copy_to_temp(src: str, suffix: str) -> str:
    """Copy src to a uniquely-named temp file so ws_server can serve it
    without exposing an arbitrary filesystem path."""
    fd, dst = tempfile.mkstemp(prefix="jarvis_preview_", suffix=suffix)
    os.close(fd)
    shutil.copyfile(src, dst)
    return dst


def _docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """Best-effort conversion of a .docx to a simple PDF. Uses python-docx
    to extract text and reportlab to render it. Returns True on success,
    False if any library is missing or the conversion fails — callers
    should fall back to a text card."""
    try:
        from docx import Document
    except Exception as e:
        print(f"[file_manager] python-docx missing: {e}")
        return False
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.pdfgen import canvas as _canvas
    except Exception as e:
        print(f"[file_manager] reportlab missing: {e}")
        return False

    try:
        doc = Document(docx_path)
        paragraphs = [p.text for p in doc.paragraphs]
        # Include table text too — rough, but captures the content.
        for tbl in doc.tables:
            for row in tbl.rows:
                paragraphs.append(" | ".join(c.text for c in row.cells))
    except Exception as e:
        print(f"[file_manager] docx read error: {e}")
        return False

    try:
        c = _canvas.Canvas(pdf_path, pagesize=letter)
        width, height = letter
        margin = 0.75 * inch
        y = height - margin
        line_height = 14
        max_chars = 95

        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin, y, Path(docx_path).name[:100])
        y -= line_height * 1.6
        c.setFont("Helvetica", 10)

        def _wrap(s: str) -> list[str]:
            words = s.split()
            if not words:
                return [""]
            lines: list[str] = []
            cur = ""
            for w in words:
                if len(cur) + len(w) + 1 > max_chars:
                    if cur:
                        lines.append(cur)
                    cur = w
                else:
                    cur = (cur + " " + w) if cur else w
            if cur:
                lines.append(cur)
            return lines

        for para in paragraphs:
            for line in _wrap(para):
                if y < margin:
                    c.showPage()
                    c.setFont("Helvetica", 10)
                    y = height - margin
                c.drawString(margin, y, line)
                y -= line_height
            y -= line_height * 0.3  # paragraph spacing

        c.save()
        return True
    except Exception as e:
        print(f"[file_manager] reportlab render error: {e}")
        return False


def prepare_preview(filepath: str) -> dict:
    """Prepare a file for display. Behavior varies by file type.

    Quick Look mode: spawn qlmanage and return a marker dict with no
    WebView payload; the caller still uses the same confirmation flow.
    """
    file_type = get_file_type(filepath)
    base: dict = {
        "file_type": file_type,
        "serve_url": None,
        "text_content": None,
        "temp_path": None,
    }

    if FILE_PREVIEW_MODE == "quicklook":
        try:
            subprocess.Popen(
                ["qlmanage", "-p", filepath],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
        except Exception as e:
            print(f"[file_manager] qlmanage failed: {e}")
        base["mode"] = "quicklook"
        return base

    base["mode"] = "orb"

    if file_type == "image":
        try:
            temp = _copy_to_temp(filepath, Path(filepath).suffix.lower() or ".img")
            _set_active_preview(temp)
            base["serve_url"] = _PREVIEW_SERVE_URL
            base["temp_path"] = temp
        except Exception as e:
            print(f"[file_manager] image copy failed: {e}")
            base["text_content"] = _other_card_text(filepath)
            base["file_type"] = "other"
        return base

    if file_type == "pdf":
        # Serve the original directly — no temp copy needed.
        _set_active_preview(filepath)
        base["serve_url"] = _PREVIEW_SERVE_URL
        return base

    if file_type == "text":
        base["text_content"] = _read_text_file(filepath)
        return base

    if file_type == "docx":
        fd, temp_pdf = tempfile.mkstemp(prefix="jarvis_preview_", suffix=".pdf")
        os.close(fd)
        if _docx_to_pdf(filepath, temp_pdf):
            _set_active_preview(temp_pdf)
            base["serve_url"] = _PREVIEW_SERVE_URL
            base["temp_path"] = temp_pdf
            base["file_type"] = "pdf"  # the orb renders it as a PDF
            return base
        # Fallback: text card with docx text content.
        try:
            os.remove(temp_pdf)
        except OSError:
            pass
        try:
            from docx import Document
            doc = Document(filepath)
            extracted = "\n".join(p.text for p in doc.paragraphs)
            base["text_content"] = (
                f"{Path(filepath).name}\n\n{extracted}"
                if extracted.strip()
                else _other_card_text(filepath)
            )
            base["file_type"] = "text"
        except Exception:
            base["text_content"] = _other_card_text(filepath)
            base["file_type"] = "other"
        return base

    # "other" → text card
    base["text_content"] = _other_card_text(filepath)
    return base


def cleanup_temp_preview(filepath: Optional[str]) -> None:
    """Delete a temp preview file created by prepare_preview. Always safe
    to call — no-op if path is None, empty, or doesn't exist. Never
    touches the user's original file."""
    _clear_active_preview()
    if not filepath:
        return
    try:
        p = Path(filepath)
        # Defensive: only delete files that look like our temp artifacts.
        if p.name.startswith("jarvis_preview_") and p.is_file():
            p.unlink()
    except Exception as e:
        print(f"[file_manager] cleanup_temp_preview error: {e}")


# ── Active preview (read by ws_server /preview_file route) ──────────────────

_active_preview_path: Optional[str] = None


def _set_active_preview(path: str) -> None:
    global _active_preview_path
    _active_preview_path = path


def _clear_active_preview() -> None:
    global _active_preview_path
    _active_preview_path = None


def get_active_preview_path() -> Optional[str]:
    """Called by ws_server to answer /preview_file requests."""
    return _active_preview_path


# ── File operations ──────────────────────────────────────────────────────────

def move_file(source_path: str, destination_path: str) -> tuple[bool, str]:
    """Move the original file to destination. destination_path may be a
    directory (in which case the original filename is preserved) or a
    full target path. Never touches temp preview files.

    Refuses any source or destination inside Jarvis's own project or
    data directory — this is the final safety net after search-level
    exclusion. Even if a protected file somehow makes it through to a
    confirmation prompt, the actual filesystem operation will not run."""
    try:
        src = Path(source_path).expanduser()
        dst = Path(destination_path).expanduser()
        if is_protected_path(str(src)):
            return (False, "that file is part of Jarvis's own system and can't be moved")
        if is_protected_path(str(dst)):
            return (False, "that destination is inside Jarvis's own system")
        if not src.exists():
            return (False, f"source file not found: {source_path}")
        if dst.is_dir():
            dst = dst / src.name
        # Don't silently clobber an existing file at the destination.
        if dst.exists():
            return (False, f"a file already exists at {dst}")
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(src), str(dst))
        return (True, str(dst))
    except Exception as e:
        return (False, f"move failed: {e}")


def rename_file(filepath: str, new_name: str) -> tuple[bool, str]:
    """Rename the file in place, keeping its directory. If new_name has no
    extension, preserve the original file's extension.

    Refuses any path inside Jarvis's own project or data directory."""
    try:
        src = Path(filepath).expanduser()
        if is_protected_path(str(src)):
            return (False, "that file is part of Jarvis's own system and can't be renamed")
        if not src.exists():
            return (False, f"file not found: {filepath}")
        new = (new_name or "").strip()
        if not new:
            return (False, "new name is empty")
        # Strip accidental path separators from spoken names.
        new = new.replace("/", "_").replace("\\", "_")
        new_path = src.with_name(new)
        if not new_path.suffix and src.suffix:
            new_path = new_path.with_suffix(src.suffix)
        if new_path == src:
            return (False, "new name matches current name")
        if new_path.exists():
            return (False, f"a file named {new_path.name} already exists")
        src.rename(new_path)
        return (True, str(new_path))
    except Exception as e:
        return (False, f"rename failed: {e}")


# ── Destination resolution ────────────────────────────────────────────────────

def resolve_destination(spoken_location: Optional[str]) -> Optional[str]:
    """Map a spoken location to a real path. Falls through unchanged for
    anything that isn't a known alias, so an explicit path from the LLM
    still works."""
    if not spoken_location:
        return None
    s = str(spoken_location).strip()
    if not s:
        return None
    key = s.lower().strip().strip(".?!,;:")
    # Allow phrases like "to the desktop" or "my desktop"
    for prefix in ("to the ", "to my ", "the ", "my ", "to "):
        if key.startswith(prefix):
            key = key[len(prefix):]
            break
    if key in _COMMON_DESTINATIONS:
        return _COMMON_DESTINATIONS[key]
    # Already an absolute or ~-expandable path — pass through.
    return str(Path(s).expanduser())
